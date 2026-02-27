"""Extract text and images from Word docs and search by keyword."""
import re
import xml.etree.ElementTree as ET
import zipfile
from io import BytesIO
from typing import List, Optional, Tuple
from docx import Document

# Word XML namespaces
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
}


def _get_text(elem) -> str:
    """Extract all text from an element (paragraph or cell)."""
    if elem is None:
        return ""
    # Handle any namespace on w:t
    texts = []
    for e in elem.iter():
        if e.tag.endswith("}t") or e.tag == "t":
            texts.append(e.text or "")
    return "".join(texts)


def _find_blip_embed(elem) -> Optional[str]:
    """Find r:embed (relationship id) for an image in drawing."""
    if elem is None:
        return None
    r_embed = "http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
    for e in elem.iter():
        if "blip" in (e.tag or "").lower():
            for k, v in (e.attrib or {}).items():
                if "embed" in k.lower() and v:
                    return v
    return None


def get_nearest_image(
    file_bytes: bytes,
    filename: str,
    matching_texts: List[str],
    images_list: List[Tuple[str, bytes, str]],
) -> Optional[Tuple[bytes, str]]:
    """
    From the document, find the image that appears closest to text matching the keyword.
    images_list: list of (doc_name, image_bytes, image_name) for this document.
    Returns (image_bytes, image_name) or None if no image / no match.
    """
    if not matching_texts or not images_list:
        return (images_list[0][1], images_list[0][2]) if images_list else None
    doc_images = [(b, n) for fn, b, n in images_list if fn == filename]
    if not doc_images:
        return None
    try:
        with zipfile.ZipFile(BytesIO(file_bytes), "r") as z:
            doc_xml = z.read("word/document.xml")
            try:
                rels_xml = z.read("word/_rels/document.xml.rels")
            except KeyError:
                return doc_images[0]
        root = ET.fromstring(doc_xml)
        body = next((c for c in root if c.tag and "body" in (c.tag or "").lower()), None)
        if body is None:
            return doc_images[0]
        # Build (position, type, value): 'text' -> text, 'image' -> rId
        blocks = []
        rid_to_target = {}
        try:
            rels_root = ET.fromstring(rels_xml)
            for rel in rels_root.findall(".//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship"):
                r_id = rel.get("Id")
                target = rel.get("Target", "")
                if "media/" in target:
                    rid_to_target[r_id] = "word/" + target.lstrip("/")
        except Exception:
            pass
        pos = 0
        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                para_text = _get_text(child).strip()
                if para_text:
                    blocks.append((pos, "text", para_text))
                    pos += 1
                embed = _find_blip_embed(child)
                if embed:
                    blocks.append((pos, "image", embed))
                    pos += 1
            elif tag == "tbl":
                for row in child:
                    if row.tag.endswith("}tr") or row.tag == "tr":
                        for cell in row:
                            if cell.tag.endswith("}tc") or cell.tag == "tc":
                                cell_text = _get_text(cell).strip()
                                if cell_text:
                                    blocks.append((pos, "text", cell_text))
                                    pos += 1
                for row in child:
                    if row.tag.endswith("}tr") or row.tag == "tr":
                        for cell in row:
                            if cell.tag.endswith("}tc") or cell.tag == "tc":
                                embed = _find_blip_embed(cell)
                                if embed:
                                    blocks.append((pos, "image", embed))
                                    pos += 1
        image_blocks = [(i, v) for i, (_, t, v) in enumerate(blocks) if t == "image"]
        if not image_blocks:
            return doc_images[0]
        matching_indices = set()
        for i, (_, typ, block_text) in enumerate(blocks):
            if typ != "text" or not block_text:
                continue
            if any(mt.strip().lower() in block_text.lower() for mt in matching_texts if mt.strip()):
                matching_indices.add(i)
        if not matching_indices:
            return doc_images[0]
        def dist(img_idx):
            return min(abs(img_idx - ti) for ti in matching_indices)
        nearest_img_idx, nearest_rid = min(image_blocks, key=lambda x: dist(x[0]))
        target_path = rid_to_target.get(nearest_rid)
        if target_path:
            with zipfile.ZipFile(BytesIO(file_bytes), "r") as z:
                try:
                    img_bytes = z.read(target_path)
                    name = target_path.split("/")[-1]
                    return (img_bytes, name)
                except Exception:
                    pass
        return doc_images[0]
    except Exception:
        return doc_images[0] if doc_images else None


def extract_images_from_docx(file_bytes: bytes, filename: str) -> List[Tuple[str, bytes, str]]:
    """
    Extract all images from a .docx file (docx is a zip with word/media/).
    Returns list of (source_doc_name, image_bytes, image_filename).
    """
    results = []
    try:
        with zipfile.ZipFile(BytesIO(file_bytes), "r") as z:
            for name in z.namelist():
                if name.startswith("word/media/"):
                    image_name = name.split("/")[-1]
                    results.append((filename, z.read(name), image_name))
    except Exception:
        pass
    return results


def extract_text_from_docx(file_bytes: bytes, filename: str) -> List[Tuple[str, str]]:
    """
    Extract all paragraphs from a .docx file.
    Returns list of (filename, paragraph_text).
    """
    results = []
    try:
        doc = Document(BytesIO(file_bytes))
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                results.append((filename, text))
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        results.append((filename, text))
    except Exception as e:
        results.append((filename, f"[Error reading file: {e}]"))
    return results


def search_keyword(paragraphs: List[Tuple[str, str]], keyword: str) -> List[Tuple[str, str]]:
    """
    Find all (filename, paragraph) where paragraph contains keyword (case-insensitive).
    """
    if not keyword or not keyword.strip():
        return []
    k = keyword.strip().lower()
    return [(fname, text) for fname, text in paragraphs if k in text.lower()]


def extract_text_from_pdf(file_bytes: bytes, filename: str) -> List[Tuple[str, str]]:
    """
    Extract text from each page of a PDF file.
    Returns list of (filename, text_block) — one block per page, split by paragraphs where possible.
    """
    results = []
    try:
        from pypdf import PdfReader
        reader = PdfReader(BytesIO(file_bytes))
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                # Split into paragraphs (double newline) or lines for easier search
                for block in text.split("\n\n"):
                    block = block.strip()
                    if block:
                        results.append((filename, block))
    except Exception as e:
        results.append((filename, f"[Error reading PDF: {e}]"))
    return results


# Delimiter used to join Excel row cells (rare in normal cell content)
EXCEL_ROW_SEP = "\x00"


def get_excel_headers(file_bytes: bytes, filename: str) -> dict:
    """
    Get the first row (column names) of each sheet. Uses actual row data
    to determine column count so all column names are fetched even when
    the first row has fewer cells than later rows.
    """
    headers = {}
    try:
        name_lower = filename.lower()
        if name_lower.endswith(".xlsx"):
            import openpyxl
            wb = openpyxl.load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
            for sheet in wb.worksheets:
                sheet_name = sheet.title
                rows = list(sheet.iter_rows(values_only=True))
                if not rows:
                    continue
                max_col = max(len(r) for r in rows)
                first_row = list(rows[0])
                first_row = first_row + [None] * (max_col - len(first_row))
                headers[(filename, sheet_name)] = [
                    str(c).strip() if c is not None else "" for c in first_row
                ]
            wb.close()
        elif name_lower.endswith(".xls"):
            import xlrd
            wb = xlrd.open_workbook(file_contents=file_bytes)
            for sheet in wb.sheets():
                sheet_name = sheet.name
                if sheet.nrows > 0 and sheet.ncols > 0:
                    headers[(filename, sheet_name)] = [
                        str(sheet.cell_value(0, col_idx)).strip()
                        for col_idx in range(sheet.ncols)
                    ]
    except Exception:
        pass
    return headers


# Row index (0-based) to use as "dates row" in vacation-tracker style sheets (e.g. row 6 = index 5)
DATES_ROW_INDEX = 5


def get_excel_dates_row(file_bytes: bytes, filename: str) -> dict:
    """
    Get the dates row (e.g. row 6) of each sheet for vacation-tracker style Excel.
    Returns dict: (filename, sheet_name) -> [cell1, cell2, ...] padded to max_col.
    Empty dict if sheet has fewer rows.
    """
    dates_row = {}
    try:
        name_lower = filename.lower()
        if name_lower.endswith(".xlsx"):
            import openpyxl
            wb = openpyxl.load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
            for sheet in wb.worksheets:
                sheet_name = sheet.title
                rows = list(sheet.iter_rows(values_only=True))
                if len(rows) <= DATES_ROW_INDEX:
                    continue
                max_col = max(len(r) for r in rows)
                row6 = list(rows[DATES_ROW_INDEX])
                row6 = row6 + [None] * (max_col - len(row6))
                dates_row[(filename, sheet_name)] = [
                    str(c).strip() if c is not None else "" for c in row6[:max_col]
                ]
            wb.close()
        elif name_lower.endswith(".xls"):
            import xlrd
            wb = xlrd.open_workbook(file_contents=file_bytes)
            for sheet in wb.sheets():
                sheet_name = sheet.name
                if sheet.nrows > DATES_ROW_INDEX and sheet.ncols > 0:
                    dates_row[(filename, sheet_name)] = [
                        str(sheet.cell_value(DATES_ROW_INDEX, col_idx)).strip()
                        for col_idx in range(sheet.ncols)
                    ]
    except Exception:
        pass
    return dates_row


def extract_text_from_excel(file_bytes: bytes, filename: str) -> List[Tuple[str, str]]:
    """
    Extract text from Excel by row. Returns one (filename, row_text) per row.
    row_text = "[Sheet: name]" + EXCEL_ROW_SEP + cell1 + EXCEL_ROW_SEP + cell2 + ...
    so we can split and display as table. Search matches if keyword is in any cell.
    """
    results = []
    try:
        name_lower = filename.lower()
        if name_lower.endswith(".xlsx"):
            import openpyxl
            wb = openpyxl.load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
            for sheet in wb.worksheets:
                sheet_name = sheet.title
                rows = list(sheet.iter_rows(values_only=True))
                max_col = max(len(r) for r in rows) if rows else 0
                for row in rows:
                    row_list = list(row) + [None] * (max_col - len(row))
                    row_cells = [str(c).strip() if c is not None else "" for c in row_list[:max_col]]
                    if any(row_cells):
                        row_text = f"[Sheet: {sheet_name}]{EXCEL_ROW_SEP}" + EXCEL_ROW_SEP.join(row_cells)
                        results.append((filename, row_text))
            wb.close()
        elif name_lower.endswith(".xls"):
            import xlrd
            wb = xlrd.open_workbook(file_contents=file_bytes)
            for sheet in wb.sheets():
                sheet_name = sheet.name
                for row_idx in range(sheet.nrows):
                    row_cells = [
                        str(sheet.cell_value(row_idx, col_idx)).strip()
                        for col_idx in range(sheet.ncols)
                    ]
                    if any(row_cells):
                        row_text = f"[Sheet: {sheet_name}]{EXCEL_ROW_SEP}" + EXCEL_ROW_SEP.join(row_cells)
                        results.append((filename, row_text))
    except Exception as e:
        results.append((filename, f"[Error reading Excel: {e}]"))
    return results


def get_word_suggestions(
    paragraphs: List[Tuple[str, str]],
    prefix: str = "",
    max_suggestions: int = 20,
) -> List[str]:
    """
    Extract unique words from all paragraphs, optionally filter by prefix.
    Returns list of words sorted by frequency (most common first).
    """
    from collections import Counter
    word_count = Counter()
    for _fname, text in paragraphs:
        if not text or not isinstance(text, str):
            continue
        words = re.findall(r"\b[a-zA-Z0-9]{2,}\b", text)
        for w in words:
            word_count[w.lower()] += 1
    if not word_count:
        return []
    prefix = prefix.strip().lower()
    if prefix:
        filtered = [(w, c) for w, c in word_count.items() if w.startswith(prefix) or prefix in w]
    else:
        filtered = list(word_count.items())
    filtered.sort(key=lambda x: (-x[1], x[0]))
    return [w for w, _ in filtered[:max_suggestions]]
